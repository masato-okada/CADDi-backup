import requests
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO


# Airtable APIの設定
API_KEY = 'patxGbJMFYSoqLlEZ.1d424338ac83c450a3560ed63444d3730fe122c58261d1effb9b58db827c05de'
BASE_ID = 'appEausi0Ru8tckzo'
TABLE_NAME = 'パトロールデータ'

# Airtable APIエンドポイント
url = f'https://api.airtable.com/v0/{BASE_ID}/{TABLE_NAME}'
headers = {
    'Authorization': f'Bearer {API_KEY}'
}

# データ取得
response = requests.get(url, headers=headers)
params = {
    'view': '報告書生成用'  # グリッドビューの名前を指定
}
response = requests.get(url, headers=headers, params=params)
if response.status_code == 200:
    data = response.json()
else:
    print(f"エラー: {response.status_code}")
    exit()

# Word報告書の作成
doc = Document()
doc.add_heading('Airtable 報告書', level=1)

# Airtableデータをレポートに追加
record_counter = 1
for record in data['records']:
    fields = record['fields']

    # レコードIDと連番を追加
    doc.add_heading(f" {record_counter} - ID: {record['id']}", level=2)
    record_counter += 1

    # テーブルを作成
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'

    # 1行目に指摘内容と是正結果を追加し、背景をグレーに設定
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '指摘内容'
    hdr_cells[1].text = '是正結果'
    for cell in hdr_cells:
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), 'D9D9D9')  # グレー色に設定
        tcPr.append(shd)    
        # フォントをMeiryo UIに設定
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Meiryo UI'
                run.font.size = Pt(10)        

    row_cells = table.rows[1].cells
    row_cells[0].text = fields.get('指摘内容', 'データなし')
    row_cells[1].text = fields.get('是正結果', 'データなし')
    for cell in row_cells:
        # フォントをMeiryo UIに設定
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Meiryo UI'
                run.font.size = Pt(10)    


    # 画像リンクを追加
    if '指摘内容(添付ファイル)' in fields:
        photos = fields['指摘内容(添付ファイル)']
        if isinstance(photos, list):
            for photo in photos:
                photo_url = photo['url']
                photo_response = requests.get(photo_url)
                if photo_response.status_code == 200:
                    # 画像をローカルに一時的に保存
                    image_stream = BytesIO(photo_response.content)
                    # 写真行を追加
                    photo_row = table.add_row().cells
                    photo_row[0].text = '是正写真'
                    paragraph = photo_row[1].paragraphs[0]
                    run = paragraph.add_run()
                    run.add_picture(image_stream, width=Inches(2))
                    # フォントをMeiryo UIに設定
                    for run in photo_row[0].paragraphs[0].runs:
                        run.font.name = 'Meiryo UI'
                        run.font.size = Pt(10)
                else:
                    photo_row = table.add_row().cells
                    photo_row[0].text = '是正写真'
                    photo_row[1].text = f'画像リンク: {photo_url} (ダウンロード失敗)'
                    # フォントをMeiryo UIに設定
                    for cell in photo_row:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.name = 'Meiryo UI'
                                run.font.size = Pt(10)

    # レコードごとにスペースを追加
    doc.add_paragraph('')

# Wordファイルとして保存
output_file = 'Airtable_Report.docx'
doc.save(output_file)
print(f"報告書を保存しました: {output_file}")
